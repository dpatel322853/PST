
"""
report_export_v2.py
Automated export to Excel, Word, and PDF for PST screening results.
"""
from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Dict, List, Optional
import math
import pandas as pd
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak

from pst_model_v2 import calculate_pst, validation_checks, summarize_scenario_result, PREthylene, scenario_templates


DISCLAIMER = (
    "SCREENING / ENGINEERING ESTIMATION ONLY. NOT FOR FINAL DESIGN WITHOUT VALIDATION "
    "against plant data, detailed design data, approved valve data, verified material properties, "
    "and approved engineering methods."
)


def _pick_profile_times(timeseries: pd.DataFrame) -> List[float]:
    tmax = float(timeseries["time_s"].max())
    candidates = [0.0, min(60.0, tmax), min(120.0, tmax), min(300.0, tmax)]
    vals = sorted(set(float(timeseries.iloc[(timeseries["time_s"] - t).abs().argmin()]["time_s"]) for t in candidates))
    return vals


def _plot_png_bytes(results_protected: Dict[str, object], results_unprotected: Optional[Dict[str, object]] = None) -> Dict[str, bytes]:
    ts = results_protected["timeseries"]
    pf = results_protected["profiles"]
    figs = {}

    def savefig_bytes(fig):
        bio = BytesIO()
        fig.savefig(bio, format="png", dpi=180, bbox_inches="tight")
        plt.close(fig)
        return bio.getvalue()

    fig, ax = plt.subplots(figsize=(8,4.5))
    ax.plot(ts["time_s"], ts["superheater_outlet_temp_c"], label="SH outlet BC")
    ax.plot(ts["time_s"], ts["section_a_exit_fluid_temp_c"], label="TT location fluid")
    ax.plot(ts["time_s"], ts["min_cs_fluid_temp_c"], label="Min CS fluid")
    if results_unprotected is not None:
        tsu = results_unprotected["timeseries"]
        ax.plot(tsu["time_s"], tsu["min_cs_fluid_temp_c"], '--', label="Min CS fluid unprotected")
    ax.set_title("Fluid temperature vs time")
    ax.set_xlabel("Time (s)")
    ax.set_ylabel("Temperature (°C)")
    ax.grid(True, alpha=0.3)
    ax.legend(fontsize=8)
    figs["fluid_vs_time"] = savefig_bytes(fig)

    fig, ax = plt.subplots(figsize=(8,4.5))
    ax.plot(ts["time_s"], ts["section_a_exit_wall_temp_c"], label="Section A exit wall")
    ax.plot(ts["time_s"], ts["min_cs_wall_temp_c"], label="Min CS wall")
    if results_unprotected is not None:
        tsu = results_unprotected["timeseries"]
        ax.plot(tsu["time_s"], tsu["min_cs_wall_temp_c"], '--', label="Min CS wall unprotected")
    ax.set_title("Metal temperature vs time")
    ax.set_xlabel("Time (s)")
    ax.set_ylabel("Temperature (°C)")
    ax.grid(True, alpha=0.3)
    ax.legend(fontsize=8)
    figs["metal_vs_time"] = savefig_bytes(fig)

    chosen = _pick_profile_times(ts)
    fig, ax = plt.subplots(figsize=(8,4.5))
    for t_sel in chosen:
        dfp = pf[(pf["time_s"] - t_sel).abs() < 1e-9]
        ax.plot(dfp["x_m"], dfp["fluid_temp_c"], label=f"t={t_sel:.0f}s")
    ax.set_title("Fluid temperature vs distance")
    ax.set_xlabel("Distance from SH outlet (m)")
    ax.set_ylabel("Fluid temperature (°C)")
    ax.grid(True, alpha=0.3)
    ax.legend(fontsize=8)
    figs["fluid_vs_distance"] = savefig_bytes(fig)

    fig, ax = plt.subplots(figsize=(8,4.5))
    for t_sel in chosen:
        dfp = pf[(pf["time_s"] - t_sel).abs() < 1e-9]
        ax.plot(dfp["x_m"], dfp["wall_temp_c"], label=f"t={t_sel:.0f}s")
    ax.set_title("Wall temperature vs distance")
    ax.set_xlabel("Distance from SH outlet (m)")
    ax.set_ylabel("Wall temperature (°C)")
    ax.grid(True, alpha=0.3)
    ax.legend(fontsize=8)
    figs["wall_vs_distance"] = savefig_bytes(fig)

    fig, ax = plt.subplots(figsize=(8,4.5))
    ax.plot(ts["time_s"], 100.0 * ts["valve_open_frac"])
    ax.set_title("Valve open fraction vs time")
    ax.set_xlabel("Time (s)")
    ax.set_ylabel("Valve open (%)")
    ax.grid(True, alpha=0.3)
    figs["valve_vs_time"] = savefig_bytes(fig)

    fig, ax = plt.subplots(figsize=(8,4.5))
    ax.plot(ts["time_s"], ts["flowrate_kg_s"])
    ax.set_title("Flowrate vs time")
    ax.set_xlabel("Time (s)")
    ax.set_ylabel("Flowrate (kg/s)")
    ax.grid(True, alpha=0.3)
    figs["flow_vs_time"] = savefig_bytes(fig)

    return figs


def export_excel_report(filepath: str,
                        results_protected: Dict[str, object],
                        results_unprotected: Optional[Dict[str, object]] = None,
                        scenario_summary: Optional[pd.DataFrame] = None) -> str:
    path = Path(filepath)
    inp = results_protected["inputs"]
    pst = calculate_pst(inp, results_protected, results_unprotected)
    val = validation_checks(inp, results_protected)
    char = results_protected["characteristics"]
    ts = results_protected["timeseries"]
    pf = results_protected["profiles"]
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        pd.DataFrame([results_protected["summary"]]).T.to_excel(writer, sheet_name="Protected_Summary")
        if results_unprotected is not None:
            pd.DataFrame([results_unprotected["summary"]]).T.to_excel(writer, sheet_name="Unprotected_Summary")
        pst.to_excel(writer, sheet_name="PST_Breakdown", index=False)
        char.to_excel(writer, sheet_name="Characteristics", index=False)
        val.to_excel(writer, sheet_name="Validation", index=False)
        ts.to_excel(writer, sheet_name="Protected_Timeseries", index=False)
        if results_unprotected is not None:
            results_unprotected["timeseries"].to_excel(writer, sheet_name="Unprotected_Timeseries", index=False)
        # Limit profile sheet size for convenience
        sample_times = _pick_profile_times(ts)
        pf_small = pf[pf["time_s"].isin(sample_times)]
        pf_small.to_excel(writer, sheet_name="Profiles_Selected", index=False)
        if scenario_summary is not None and len(scenario_summary) > 0:
            scenario_summary.to_excel(writer, sheet_name="Scenario_Manager", index=False)
        # assumptions sheet
        assumptions = pd.DataFrame({
            "Assumption": [
                DISCLAIMER,
                "PR EOS used for pressure-dependent density and non-ideal Cp screening adjustment.",
                "Transport properties are screening approximations; viscosity and thermal conductivity are temperature-dependent only unless user extends the model.",
                "No axial conduction, no radiation, no fittings/support thermal masses, no pressure-drop dynamics, no phase change in pipe model.",
                "Available PST is interpreted as unprotected time from initiating event to first CS wall segment reaching MDMT.",
            ]
        })
        assumptions.to_excel(writer, sheet_name="Assumptions", index=False)
    return str(path)


def _add_df_table_docx(doc: Document, df: pd.DataFrame, max_rows: int = 20):
    if df is None or df.empty:
        doc.add_paragraph("No data.")
        return
    dfi = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(dfi.columns))
    table.style = "Light List"
    hdr = table.rows[0].cells
    for i, c in enumerate(dfi.columns):
        hdr[i].text = str(c)
    for _, row in dfi.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if pd.isna(val) else f"{val}"


def export_word_report(filepath: str,
                       results_protected: Dict[str, object],
                       results_unprotected: Optional[Dict[str, object]] = None,
                       scenario_summary: Optional[pd.DataFrame] = None) -> str:
    path = Path(filepath)
    inp = results_protected["inputs"]
    pst = calculate_pst(inp, results_protected, results_unprotected)
    val = validation_checks(inp, results_protected)
    charts = _plot_png_bytes(results_protected, results_unprotected)

    doc = Document()
    doc.add_heading("LP Ethylene Superheater PST Screening Report", 0)
    doc.add_paragraph(DISCLAIMER)
    doc.add_heading("1. Executive Summary", level=1)
    summary_df = pd.DataFrame([summarize_scenario_result(inp, results_protected, results_unprotected, "Selected Case")])
    _add_df_table_docx(doc, summary_df, max_rows=10)
    doc.add_heading("2. PST Breakdown", level=1)
    _add_df_table_docx(doc, pst, max_rows=20)
    doc.add_heading("3. Characteristic Values", level=1)
    _add_df_table_docx(doc, results_protected["characteristics"], max_rows=10)
    doc.add_heading("4. Validation / Reasonableness Checks", level=1)
    _add_df_table_docx(doc, val, max_rows=20)
    if scenario_summary is not None and len(scenario_summary) > 0:
        doc.add_heading("5. Scenario Manager Summary", level=1)
        _add_df_table_docx(doc, scenario_summary, max_rows=30)
    doc.add_heading("6. Plots", level=1)
    for key, title in [
        ("fluid_vs_time", "Fluid temperature vs time"),
        ("metal_vs_time", "Metal temperature vs time"),
        ("fluid_vs_distance", "Fluid temperature vs distance"),
        ("wall_vs_distance", "Wall temperature vs distance"),
        ("valve_vs_time", "Valve position vs time"),
        ("flow_vs_time", "Flowrate vs time"),
    ]:
        doc.add_heading(title, level=2)
        img_path = path.with_name(f"_{key}.png")
        img_path.write_bytes(charts[key])
        doc.add_picture(str(img_path), width=Inches(6.5))
    doc.add_heading("7. Limitations / Not for final design without validation", level=1)
    for bullet in [
        DISCLAIMER,
        "Pressure-dependent density and non-ideal Cp are estimated with Peng-Robinson EOS for pure ethylene screening. Transport properties remain simplified.",
        "Actual exchanger dynamics, line pressure, transport properties, valve Cv-vs-travel, and sensor thermowell behavior can materially alter the result.",
        "Use plant data / validated models before relying on this report for safety decisions.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")
    doc.save(path)
    return str(path)


def _df_to_table_data(df: pd.DataFrame, max_rows: int = 25):
    if df is None or df.empty:
        return [["No data"]]
    dfi = df.head(max_rows).copy()
    data = [list(dfi.columns)]
    for _, row in dfi.iterrows():
        data.append(["" if pd.isna(v) else str(v) for v in row.tolist()])
    return data


def export_pdf_report(filepath: str,
                      results_protected: Dict[str, object],
                      results_unprotected: Optional[Dict[str, object]] = None,
                      scenario_summary: Optional[pd.DataFrame] = None) -> str:
    path = Path(filepath)
    inp = results_protected["inputs"]
    pst = calculate_pst(inp, results_protected, results_unprotected)
    val = validation_checks(inp, results_protected)
    charts = _plot_png_bytes(results_protected, results_unprotected)

    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=1.2*cm, leftMargin=1.2*cm, topMargin=1.0*cm, bottomMargin=1.0*cm)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph("LP Ethylene Superheater PST Screening Report", styles["Title"]))
    story.append(Spacer(1, 0.2*cm))
    story.append(Paragraph(DISCLAIMER, styles["BodyText"]))
    story.append(Spacer(1, 0.3*cm))

    for title, df in [
        ("Executive Summary", pd.DataFrame([summarize_scenario_result(inp, results_protected, results_unprotected, "Selected Case")])),
        ("PST Breakdown", pst),
        ("Characteristic Values", results_protected["characteristics"]),
        ("Validation / Reasonableness Checks", val),
    ]:
        story.append(Paragraph(title, styles["Heading1"]))
        tb = Table(_df_to_table_data(df, max_rows=25), repeatRows=1)
        tb.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
            ("GRID", (0,0), (-1,-1), 0.25, colors.grey),
            ("FONTSIZE", (0,0), (-1,-1), 7),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
        ]))
        story.append(tb)
        story.append(Spacer(1, 0.2*cm))

    if scenario_summary is not None and len(scenario_summary) > 0:
        story.append(Paragraph("Scenario Manager Summary", styles["Heading1"]))
        tb = Table(_df_to_table_data(scenario_summary, max_rows=30), repeatRows=1)
        tb.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
            ("GRID", (0,0), (-1,-1), 0.25, colors.grey),
            ("FONTSIZE", (0,0), (-1,-1), 7),
        ]))
        story.append(tb)
        story.append(PageBreak())

    story.append(Paragraph("Plots", styles["Heading1"]))
    for key, title in [
        ("fluid_vs_time", "Fluid temperature vs time"),
        ("metal_vs_time", "Metal temperature vs time"),
        ("fluid_vs_distance", "Fluid temperature vs distance"),
        ("wall_vs_distance", "Wall temperature vs distance"),
        ("valve_vs_time", "Valve position vs time"),
        ("flow_vs_time", "Flowrate vs time"),
    ]:
        img_path = path.with_name(f"_{key}.png")
        img_path.write_bytes(charts[key])
        story.append(Paragraph(title, styles["Heading2"]))
        story.append(Image(str(img_path), width=16*cm, height=9*cm))
        story.append(Spacer(1, 0.2*cm))
    story.append(Paragraph("Limitations / Not for final design without validation", styles["Heading1"]))
    story.append(Paragraph(DISCLAIMER, styles["BodyText"]))
    doc.build(story)
    return str(path)
